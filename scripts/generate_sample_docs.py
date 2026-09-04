import os
import sys

# Generate PDF and DOCX test files
try:
    from pypdf import PdfWriter
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    pass

try:
    import docx
except ImportError:
    pass

def create_sample_files():
    data_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'data'))
    
    sample_text = """STANDARD RESIDENTIAL LEASE AGREEMENT

PARTIES: Sterling Real Estate Management ("Landlord") and David Miller ("Tenant").

PREMISES: 405 Grand Avenue, Apartment 12B.

TERM AND RENT:
The lease term is twelve (12) months starting October 1, 2024. Monthly base rent is $2,200.00.

SECURITY DEPOSIT REQUIREMENT:
Tenant agrees to pay a security deposit of $7,700.00 (equal to 3.5 months' monthly rent) prior to taking possession.

TERMINATION NOTICE PERIOD:
Either party may terminate or non-renew this lease upon expiration by serving fifteen (15) days advance written notice.

AUTOMATIC RENEWAL:
Clause 18: Upon expiration of the initial term, this lease shall automatically renew for an additional full term of 12 months at a 10% rent increase, without requiring any notice or re-execution by either party.

MAINTENANCE RESPONSIBILITIES:
Landlord agrees to maintain central heating, plumbing, structural framework, and HVAC systems. Tenant agrees to maintain internal cleanliness.

DEPOSIT REFUND TIMELINE:
Landlord agrees to return the security deposit within twenty (21) calendar days following move-out inspection.

Dated: September 1, 2024.
Landlord: Sterling Real Estate Management
Tenant: David Miller
"""

    # 1. Save TXT file
    txt_path = os.path.join(data_dir, "sample_lease_test.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(sample_text)
    print(f"Created TXT sample file: {txt_path}")

    # 2. Save MD file
    md_path = os.path.join(data_dir, "sample_lease_test.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# Residential Lease Agreement\n\n" + sample_text)
    print(f"Created MD sample file: {md_path}")

    # 3. Create DOCX file using python-docx
    try:
        doc_path = os.path.join(data_dir, "sample_lease_test.docx")
        doc = docx.Document()
        doc.add_heading("STANDARD RESIDENTIAL LEASE AGREEMENT", level=1)
        for paragraph in sample_text.split("\n\n"):
            doc.add_paragraph(paragraph)
        doc.save(doc_path)
        print(f"Created DOCX sample file: {doc_path}")
    except Exception as e:
        print("Could not create DOCX file:", e)

    # 4. Create PDF file using fpdf / reportlab or pypdf canvas
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        for line in sample_text.split("\n"):
            pdf.cell(200, 7, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
        pdf_path = os.path.join(data_dir, "sample_lease_test.pdf")
        pdf.output(pdf_path)
        print(f"Created PDF sample file: {pdf_path}")
    except ImportError:
        # Install fpdf2 if needed or use reportlab
        print("fpdf not available, installing fpdf2...")
        os.system("pip install fpdf2")
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        for line in sample_text.split("\n"):
            pdf.cell(200, 7, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
        pdf_path = os.path.join(data_dir, "sample_lease_test.pdf")
        pdf.output(pdf_path)
        print(f"Created PDF sample file: {pdf_path}")

if __name__ == "__main__":
    create_sample_files()
